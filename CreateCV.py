from docx import Document
from json import loads
from Model.forms import (formacao, emprego, habilidades,
						projetos, jsonfields, pessoais)

class CVWritter():
	def __init__(self, data, CVmodel):
		self.data = data
		self.CVmodel = CVmodel

	def createCV(self):
		self.escrever_dados_pessoais()
		self.escrever_formacao_academica()
		self.escrever_experiencia_profisional()
		self.escrever_habilidades()
		self.escrever_projetos()

	#Dados Pessoais
	def escrever_dados_pessoais(self):
		self.CVmodel.paragraphs[0].text = self.data['pessoais']['nome']
		self.CVmodel.paragraphs[1].text = ''.join([self.data['area de interesse'][f'{x+1}']+' | ' for x in range(len(self.data['area de interesse']))])[:-3]

		pessoais = [self.data['pessoais']['cidade_estado'], self.data['pessoais']['telefone'], 
					self.data['pessoais']['email'], self.data['pessoais']['linkedin'], self.data['pessoais']['github']]
		for i in range(2, 7):
			self.CVmodel.paragraphs[i].text += pessoais[i-2]
		
	#Formação academica
	def escrever_formacao_academica(self):
		str_formacao = ''
		for x in range(len(self.data['formacao'])):
			x = str(x+1)
			self.data_using = self.data['formacao']

			str_formacao += f"Curso: {self.data_using[x]['curso']}\n"
			str_formacao += f"{self.data_using[x]['instituicao']} - {self.data_using[x]['cidade']} - {self.data_using[x]['ano']}"
			if self.data_using[x]['horas']:
				str_formacao += f" - {self.data_using[x]['horas']} horas\n"
			else:
				str_formacao += '\n'
		self.CVmodel.paragraphs[9].text = str_formacao[:-1]


	#EXPERIÊNCIA PROFISSIONAL
	def escrever_experiencia_profisional(self):
		str_experiencia = ''
		for x  in range(len(self.data['empregos'])):
			x = str(x+1)
			self.data_using = self.data['empregos']
			str_experiencia += f"{self.data_using[x]['empresa']} - {self.data_using[x]['cargo']}\nInicio: {self.data_using[x]['entrada']}"
			
			if 'dias' in self.data_using[x]['saida'].split():
				str_experiencia += f' |  Até os dias atuais\n'
			else:
				str_experiencia += f" | Término: {self.data_using[x]['saida']}\n"

			str_experiencia += f"Responsabilidades: {self.data_using[x]['resumo']}\n\n"

		self.CVmodel.paragraphs[12].text = str_experiencia[:-2]

	#HABILIDADES
	def escrever_habilidades(self):
		str_habilidades = ''
		for x in range(len(self.data['habilidades'])):
			self.data_using = self.data['habilidades']
			x = str(x+1)
			str_habilidades += f"{self.data_using[x]['tecnologia']} - {self.data_using[x]['nivel']}\n"
		self.CVmodel.paragraphs[15].text = str_habilidades

	#PROJETOS DESENVOLVIDOS
	def escrever_projetos(self):
		str_projetos = ''
		for x in range(len(self.data['projetos'])):
			self.data_using = self.data['projetos']
			x = str(x+1)

			str_projetos += f"Nome: {self.data_using[x]['nome']}\nTecnologias: {self.data_using[x]['tecnologias']}\n"
			str_projetos += f"Resumo: {self.data_using[x]['resumo']}\n\n"
		self.CVmodel.paragraphs[18].text = str_projetos[:-2]

	def save(self):
		filename = f"Files/{self.data['pessoais']['nome']}-CV.docx"
		self.CVmodel.save(filename)
		return filename


	def __repr__(self):
		return f"< Curriculum - {self.data['pessoais']['nome']} >"


if __name__ == '__main__':
	data = loads(jsonfields)
	model = Document('Model/modelo.docx')
	Pedro = CVWritter(data, model)
	Pedro.createCV()
	caminho = Pedro.save()

