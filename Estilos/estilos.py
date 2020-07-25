from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx import Document

class Padrao():
	def __init__(self, CV):
		self.CV = CV

	def subtitles(text, document):
		...
	
	def pessoais(text, document):
		...

	def formacao(text, document):
		...

	def experiencia(text, document):
		...

	def habilidades(text, document):
		...
		
	def projetos(text, document):
		...

	def save():
		...

filename = '../Files/José Pedro da Silva Gomes-CV.docx'
CV = Document(filename)
base = Document()

def subtitle(text, document):

	for i in range(len(document.paragraphs)):
		print(f'{i} - {text[i]}')
		
		if i == 0:
			document.paragraphs[i].paragraph_format.space_before = Pt(1)
			document.paragraphs[i].paragraph_format.space_after = Pt(1)
			document.paragraphs[i].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
			run = document.paragraphs[i].add_run(text[i])
			run.font.name = 'Calibri'
			run.font.size = Pt(20)

		elif i == 1:
			document.paragraphs[i].paragraph_format.space_before = Pt(1)
			document.paragraphs[i].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
			run = document.paragraphs[i].add_run(text[i])
			run.font.name = 'Calibri'
			run.font.size = Pt(13)
			run.font.color.rgb = RGBColor(0x00, 0xBF, 0xFF)

		elif text[i][-3:] == '___':

			run_ = document.paragraphs[i].add_run(text[i] + '______________')
			run_.font.color.rgb = RGBColor(0x41, 0x69, 0xE1)
			run_.font.name = 'Calibri'

			document.paragraphs[i+1].paragraph_format.left_indent = Inches(0.5)
			run = document.paragraphs[i+1].add_run(text[i+1])
			run.font.name = 'Calibri'
			run.font.size = Pt(13)
			run.font.color.rgb = RGBColor(0x41, 0x69, 0xE1)

def pessoais(text, document):

	for i in range(len(document.paragraphs)):
		print(f'{i} - {text[i]}')
		
		if i >= 2 and i <= 6:
			document.paragraphs[i].paragraph_format.space_before = Pt(1)
			document.paragraphs[i].paragraph_format.space_after = Pt(1)
			run1 = document.paragraphs[i].add_run(text[i].split(':')[0]+': ')
			run1.font.name = 'Calibri'
			run1.font.size = Pt(11)
			run1.font.bold = True

			run2 = document.paragraphs[i].add_run(text[i].split(':')[1])
			run2.font.name = 'Calibri'
			run2.font.size = Pt(11)

def formacao(text, document):
	...

def experiencia(text, document):
	...

def habilidades(text, document):
	...

def projetos(text, document):
	...

paragraphs = []
for paragraph in CV.paragraphs:
	for line in paragraph.text.split('\n'):
		paragraphs.append(line)

for paragraph in paragraphs:
	base.add_paragraph()

subtitle(paragraphs, base)
pessoais(paragraphs, base)

base.save("../Files/base.docx")


